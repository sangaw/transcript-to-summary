from __future__ import annotations

import argparse
import os
from glob import glob
from typing import Optional

from .logger import get_logger
from .reader import read_document
from .summarizer import summarize_text, summarize_document

logger = get_logger(__name__)


def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Read a document and provide a summary.")
    parser.add_argument("--input", required=False, help="Path to input file (.txt, .docx, .pdf)")
    parser.add_argument("--sentences", type=int, default=3, help="Number of sentences in summary")
    parser.add_argument(
        "--generate-sample-docx",
        action="store_true",
        help="If the input .docx does not exist, create a sample transcript there before summarizing.",
    )
    parser.add_argument(
        "--use-data-folders",
        action="store_true",
        help="Process all .docx in data/input and write summaries to data/output.",
    )
    return parser.parse_args(argv)


def _create_sample_docx(path: str) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("python-docx is required to generate a sample .docx. Install and retry.") from exc

    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    doc = Document()
    for line in [
        "Speaker 1: Hello everyone, and welcome to today's meeting.",
        "[00:00:15] Speaker 2: Let's discuss the new project proposal.",
        "Speaker 1: Main points include expanding market reach and optimizing workflows.",
        "Speaker 3: What about the budget allocation for this expansion?",
        "Speaker 2: Significant portion for marketing and product development.",
        "Speaker 1: Goal is a 20% increase in engagement in two quarters.",
    ]:
        doc.add_paragraph(line)
    doc.save(path)
    logger.info("Created sample DOCX at %s", path)


def _write_summary_docx(output_path: str, summary_text: str) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("python-docx is required to write .docx summaries. Install and retry.") from exc
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc = Document()
    for line in summary_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)
    logger.info("Wrote summary DOCX: %s", output_path)


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)
    try:
        if args.use_data_folders:
            input_dir = os.path.join("data", "input")
            output_dir = os.path.join("data", "output")
            os.makedirs(input_dir, exist_ok=True)
            os.makedirs(output_dir, exist_ok=True)

            docx_files = sorted(glob(os.path.join(input_dir, "*.docx")))
            if not docx_files:
                logger.info("No .docx files found in %s", input_dir)
                print(f"No .docx files found in {input_dir}")
                return 0

            for docx_path in docx_files:
                try:
                    logger.info("Processing DOCX: %s", docx_path)
                    summary = summarize_document(docx_path, sentences_count=max(1, args.sentences))
                    base_name = os.path.splitext(os.path.basename(docx_path))[0]
                    out_path = os.path.join(output_dir, f"{base_name}_summary.docx")
                    _write_summary_docx(out_path, summary)
                except Exception as file_exc:
                    logger.exception("Failed to process %s: %s", docx_path, file_exc)
            print(f"Processed {len(docx_files)} file(s). Summaries are in {output_dir}")
            return 0

        if not args.input:
            raise SystemExit("--input is required unless --use-data-folders is provided")

        extension = os.path.splitext(args.input)[1].lower()
        if extension == ".docx":
            if not os.path.exists(args.input) and args.generate_sample_docx:
                _create_sample_docx(args.input)
            if not os.path.exists(args.input):
                raise FileNotFoundError(
                    f"DOCX not found: {args.input}. Use --generate-sample-docx to create a sample here."
                )
            result = summarize_document(args.input, sentences_count=max(1, args.sentences))
            print(result)
        else:
            if not os.path.exists(args.input):
                raise FileNotFoundError(f"Input not found: {args.input}")
            text = read_document(args.input)
            result = summarize_text(text, sentences_count=max(1, args.sentences))
            print(result)
        return 0
    except Exception as exc:
        logger.exception("Failed to summarize: %s", exc)
        return 1
