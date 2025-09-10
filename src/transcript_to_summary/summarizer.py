import docx
import nltk
import re
from typing import List
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.utils import get_stop_words

from .logger import get_logger

logger = get_logger(__name__)


def _ensure_punkt() -> None:
    resources = [
        ("punkt", "tokenizers/punkt"),
        ("punkt_tab", "tokenizers/punkt_tab"),
    ]
    for resource_name, resource_path in resources:
        try:
            nltk.data.find(resource_path)
        except Exception:
            try:
                nltk.download(resource_name, quiet=True)
                logger.info("Downloaded NLTK resource: %s", resource_name)
            except Exception as exc:
                logger.warning("Failed to download NLTK resource %s: %s", resource_name, exc)


def read_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    full_text: List[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    return "\n".join(full_text)


def preprocess_text(text: str) -> str:
    text = re.sub(r'\[\d{2}:\d{2}:\d{2}\]\s*', '', text)
    text = re.sub(r'(Speaker \d+|(?:[A-Z][a-z]+(?:\s[A-Z][a-z]+)*):)\s*', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def summarize_text(text: str, sentences_count: int = 5, language: str = "english") -> str:
    _ensure_punkt()
    processed_text = preprocess_text(text)
    if not processed_text:
        return ""
    parser = PlaintextParser.from_string(processed_text, Tokenizer(language))
    summarizer = TextRankSummarizer()
    summarizer.stop_words = get_stop_words(language)
    summary_sentences: List[str] = [str(sentence) for sentence in summarizer(parser.document, max(1, sentences_count))]
    return "\n".join(summary_sentences)


def summarize_document(docx_path: str, sentences_count: int = 5, language: str = "english") -> str:
    logger.info("Reading DOCX for summarization: %s", docx_path)
    raw_text = read_docx(docx_path)
    if not raw_text:
        return "Could not read any text from the document or document is empty."
    logger.info("Preprocessing text for summarization")
    return summarize_text(raw_text, sentences_count=sentences_count, language=language)


if __name__ == "__main__":
    # --- Create a dummy Word document for testing ---
    # This part helps you test the code without manually creating a .docx file
    test_doc_path = "sample_transcript.docx"
    test_content = [
        "Speaker 1: Hello everyone, and welcome to today's meeting. We have a lot to cover.",
        "[00:00:15] Speaker 2: Yes, I'm eager to discuss the new project proposal. I believe it has great potential.",
        "Speaker 1: Absolutely. The main points of the proposal include expanding our market reach and optimizing current workflows.",
        "[00:00:45] Speaker 3: What about the budget allocation for this expansion? That's a critical factor we need to consider.",
        "Speaker 2: We've allocated a significant portion for marketing and product development, ensuring a robust launch.",
        "Speaker 1: Our goal is to achieve a 20% increase in customer engagement within the next two quarters.",
        "Speaker 3: And how will we measure that engagement? Do we have specific KPIs in mind?",
        "Speaker 1: We will use a combination of website analytics, social media metrics, and direct customer feedback surveys to track our progress.",
        "This meeting concludes with a clear action plan for the marketing team to finalize the campaign details by next Friday."
    ]

    doc = docx.Document()
    for line in test_content:
        doc.add_paragraph(line)
    doc.save(test_doc_path)
    logger.info(f"Created a sample Word document: {test_doc_path}")
    # ----------------------------------------------------

    # --- Example Usage ---
    document_path = test_doc_path # Use the path to your actual transcript .docx file
    summary_length = 3 # Number of sentences in the summary

    summary_text = summarize_document(document_path, sentences_count=summary_length)

    logger.info("\n--- Generated Summary ---")
    logger.info(summary_text)

    # Clean up the test document
    import os
    os.remove(test_doc_path)
    logger.info(f"\nRemoved sample document: {test_doc_path}")